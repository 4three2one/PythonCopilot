import json
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


def process_section(paragraphs, section, attr_title, attr_text, attr_key):
    section_data = section.get('label')
    if section_data:
        get_para(paragraphs, section_data, attr_text=attr_title)

    content = section.get('content')
    pic_list = section.get('pic')
    table_list = section.get('table')
    if content:
        for item in content:
            pic_pattern = r'^pic\d+$'  # 以'pic'开始，后跟一个或多个数字，并以这些数字结束
            table_pattern = r'^table\d+$'  # 以'table'开始，后跟一个或多个数字，并以这些数字结束

            # 使用re.match检查字符串是否符合正则表达式模式
            if re.match(pic_pattern, item):
                print("Matches 'pic+int' format")
                get_pic(paragraphs, pic_list[item], attr_pic)
            elif re.match(table_pattern, item):
                print("Matches 'table+int' format")
                tittle = table_list[item]['name']
                get_para(paragraphs, tittle, attr_text=attr_pic)
                # 创建表格

                create_table(paragraphs, table_list[item]['content'], heights=table_list[item]['heights'],
                             widths=table_list[item]['widths'],
                             header_font_props=header_font_props,
                             content_font_props=content_font_props)
                paragraph = paragraphs.add_paragraph(style=None)
            else:
                get_para(paragraphs, item, attr_text=attr_text, attr_key=attr_key)

    paragraph = paragraphs.add_paragraph(style=None)
    # 递归处理嵌套数据
    nested_data = section.get('data')
    return nested_data


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_run_properties(run, font, size, RGB, bold, **kwargs):
    """
    设置Run对象的字体属性，并根据需要处理对齐方式。
    其他未指定的参数会被忽略。
    """
    # 设置基本的字体属性
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*RGB)
    run.bold = bold


def get_attr(attrtmp):
    # 设置默认值
    default_attr = {
        'font': '仿宋',
        'size': 14,
        'RGB': (0, 0, 0),
        'bold': False,
        'alignment': 'left',
        'line_spacing': 1.25,
        'space_before': 0,
        'space_after': 0,
        'indent': 0
    }

    # 使用.get方法更新字典中的值，如果键不存在，则使用默认值
    attr = {key: attrtmp.get(key, default_attr[key]) for key in default_attr}

    return attr


def get_para(paragraphs, text, attr_text=None, attr_key=None):
    para = paragraphs.add_paragraph(style=None)
    # 设置attr_text和attr_key的默认值
    attr_text = get_attr(attr_text if attr_text is not None else {})
    attr_key = get_attr(attr_key if attr_key is not None else {})

    alignment_map = {
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT
    }

    # 直接在文本中处理<strong>标签
    parts = re.split(r'(<strong>.*?</strong>)', text)
    for part in parts:
        if part.startswith('<strong>') and part.endswith('</strong>'):
            content1 = part[8:-9]
            run = para.add_run(content1)  # 高亮关键词
            set_run_properties(run, **attr_key)
        else:
            run = para.add_run(part)
            set_run_properties(run, **attr_text)

    # 设置对齐方式
    para.paragraph_format.alignment = alignment_map[attr_text['alignment']]

    # 设置行距和段间距
    para.paragraph_format.line_spacing = attr_text['line_spacing']
    para.paragraph_format.space_before = Pt(attr_text['space_before'])
    para.paragraph_format.space_after = Pt(attr_text['space_after'])

    # 处理首行缩进
    para.paragraph_format.first_line_indent = Pt(attr_text['indent'] * 12)  # 假设每个全角字符宽度约为12pt


def paragraph_center(para, text, RGB_t=(0, 0, 0), keyword='', RGB_k=(0, 0, 0), size=22, bold=True, font='仿宋',
                     indent='0', center=0, right=0, line_spacing=1.0, space_before=0, space_after=0):
    """
    在段落中添加文本，重点部分加粗变色，其他文本的颜色、字体、大小可调。
    para是段落对象，text为段落文本，keyword为关键词，RGB_k与_t对应前两个文本颜色，
    size是字号，bold给关键词加粗，font是字体，indent是首行缩进的字符数量（字符串），
    center为1表示居中，line_spacing为行距，space_before为段前距，space_after为段后距。
    """
    # 处理关键词高亮
    if keyword:
        parts = text.split(keyword)
        for i, part in enumerate(parts):
            print(i, part)
            if i > 0:
                run = para.add_run(keyword)
                set_run_properties(run, font, size, RGB_k, True)
            run = para.add_run(part)
            set_run_properties(run, font, size, RGB_t)
    else:
        run = para.add_run(text)
        set_run_properties(run, font, size, RGB_t, bold)

    # 设置首行缩进
    try:
        indent_chars = int(indent)  # 将字符串转换为整数
    except ValueError:
        indent_chars = 0  # 如果转换失败，使用默认值0

    char_width_pt = 12  # 假设每个全角字符宽度约为12pt
    total_indent_pt = Pt(indent_chars * char_width_pt)
    para.paragraph_format.first_line_indent = total_indent_pt

    # 设置对齐方式
    if center:
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif right:
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 设置行距和段间距
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)


def cell_font(cell, font_props):
    """
    设置单元格内文本的字体、大小、加粗属性，以及文本的对齐方式
    """
    alignment_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT
    }

    # 从font_props中获取对齐方式，默认为左对齐
    alignment = alignment_map.get(font_props.get('alignment', 'left'))

    for paragraph in cell.paragraphs:
        # 设置段落对齐方式
        paragraph.alignment = alignment

        # 获取或创建Run对象
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

        # 设置字体
        run.font.name = font_props.get('font', 'Arial')  # 默认字体Arial
        # run.font.name = font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_props.get('font', 'Arial'))  # 默认字体Arial
        # 设置字体大小
        run.font.size = Pt(font_props.get('size', 10))  # 默认大小10pt
        # 设置加粗
        run.bold = font_props.get('bold', False)  # 默认不加粗


def get_pic(d, attr, attr_pic):
    name = attr.get('name')
    url = attr.get('url')
    size = attr.get('size')

    # 添加包含图片的段落
    para = d.add_paragraph(style=None)
    run = para.add_run()
    run.add_picture(url, width=Cm(size))

    # 设置段落对齐方式为居中，使图片居中
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加图片说明文字的段落，并设置为居中

    get_para(d, name, attr_text=attr_pic)


def create_table(document, table_content, widths=None, heights=None, header_font_props=None, content_font_props=None,
                 style='Table Grid'):
    lengths = [len(value) for value in table_content.values()]
    rows = lengths[0] if all(length == lengths[0] for length in lengths) else "Error: Inconsistent lengths"
    cols = len(table_content.keys())

    if isinstance(rows, str):  # 检查rows是否为错误消息
        print(rows)  # 打印错误消息
        return

    table = document.add_table(rows=rows + 1, cols=cols)  # +1是因为还需要一行用于表头
    table.style = style
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 表格居中

    # 设置列宽，如果widths被提供
    if widths is not None:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):  # 防止索引越界
                    row.cells[idx].width = Cm(width)

    # 设置行高，如果heights被提供
    if heights is not None:
        for idx, height in enumerate(heights):
            if idx < len(table.rows):  # 防止索引越界
                table.rows[idx].height = Cm(height)

    # 填充表头
    header_cells = table.rows[0].cells
    for idx, key in enumerate(table_content.keys()):
        header_cells[idx].text = key
        if header_font_props is not None:
            cell_font(header_cells[idx], header_font_props)

    # 填充表格内容
    for row_idx, key in enumerate(table_content.keys()):
        for col_idx in range(1, rows + 1):  # 跳过表头，从第二行开始
            cell = table.cell(col_idx, row_idx)
            cell.text = str(table_content[key][col_idx - 1])  # col_idx-1是因为列表索引从0开始
            if content_font_props is not None:
                cell_font(cell, content_font_props)


"""
初始化数据
nPic：图片编号
nTab：表格编号
data：内容
"""
nPic = 1
nTab = 1

with open(r'./data/report.json', 'r', encoding='utf-8') as file:
    data = json.load(file)
doc_path = f"../data/test.docx"
currentTime = '2024040800'
"""
标准输入
"""
d = Document()
# 页边距
section = d.sections[-1]
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)
# 页脚
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph_run = paragraph.add_run()
"""
起始内容
"""
initParagraph = d.add_paragraph(style=None)  # 增加一个段落
paragraph_center(initParagraph, f'台风洪涝灾害风险预报信息简报', RGB_t=(43, 86, 154), size=22, bold=True,
                 font='仿宋', center=1)

initParagraph = d.add_paragraph(style=None)  # 增加一个段落
paragraph_center(initParagraph, datetime.strptime(str(currentTime), "%Y%m%d%H").strftime('%Y年%m月%d日%H时'), size=14,
                 bold=False, font='仿宋', center=1)

"""
常规内容
"""
attr_title1 = data["style"]["title-1"]
attr_title2 = data["style"]["title-2"]
attr_title3 = data["style"]["title-3"]
attr_title4 = data["style"]["title-4"]
attr_content = data["style"]["content"]
attr_pic = data["style"]["pic"]
attr_key = data["style"]["key"]
# 表头和内容的字体属性
header_font_props = data["style"]["table"]
content_font_props = data["style"]["header"]

for title1 in data:
    print(title1)
    # title1_data = t1["title-1"]
    title1_data = process_section(d, title1["title-1"], attr_title1, attr_content, attr_key)

    if not title1_data:
        continue
    for title2 in title1_data:
        title2_data = process_section(d, title2["title-2"], attr_title2, attr_content, attr_key)
        if not title2_data:
            # print(111)
            continue
        for title3 in title2_data:
            title3_data = process_section(d, title3["title-3"], attr_title3, attr_content, attr_key)
            if not title3_data:
                # print(2222)
                continue
            for title4 in title3_data:
                title4_data = process_section(d, title4["title-4"], attr_title4, attr_content, attr_key)
                if not title4_data:
                    print(3333)
                    continue

"""
终止内容
"""
table_temp = d.add_table(rows=1, cols=1)
set_cell_border(table_temp.rows[0].cells[0], top={"sz": 20, "val": "single", "color": "#2B569A", "space": "0"},
                bottom={"sz": 5, "val": "single", "color": "black", "space": "0"})
report_now = datetime.strptime(str(currentTime), "%Y%m%d%H").strftime('%Y年%m月%d日%H时')
text = f"接收单位：应急管理局\n发送单位：浙江大学韧性城市研究中心\n发送时间：{report_now}"
# 保存文档
d.save(doc_path)
